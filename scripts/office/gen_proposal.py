"""
Generate Proposal.docx (+ Proposal.dotx template) — Alio Analytics.

SALES-1 structure lens applied — cover, then structured fill-in sections in
the standard B2B proposal order: About Alio, Understanding/Scope, Approach,
Deliverables, Timeline, Pricing (Kwanza table), Terms, Signature/Acceptance.
FE-1 brand-fidelity lens: logo, brand colors, Sora/Instrument Sans/JetBrains
Mono, gradient rule, fiscal footer.
"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION

from brand import (
    FONT_DISPLAY, FONT_BODY, FONT_MONO,
    PRIMARY_700, PRIMARY_950, ACCENT_600, NEUTRAL_950, NEUTRAL_700, NEUTRAL_500, NEUTRAL_200, NEUTRAL_50, WHITE,
    COMPANY_NAME, NIF, ADDRESS, IBAN, BANK_NAME, EMAIL, WEBSITE, PHONE, TAGLINE_EN,
    LOGO_ICON_WHITE, LOGO_ICON_NAVY, OUT_DIR,
)
from docx_helpers import (
    set_default_fonts, set_run_font, add_gradient_bar, add_placeholder_line,
    set_margins, set_cell_background, strip_table_borders,
)
from to_template import docx_to_dotx

PAGE_WIDTH_CM = 21.0 - 2.5 - 2.5


def add_page_break(doc):
    doc.add_page_break()


def section_heading(doc, number, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"{number}. {title}")
    set_run_font(r, family=FONT_DISPLAY, size=18, bold=True, color=PRIMARY_700)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(10)
    add_gradient_bar(doc, 2.4, height_cm=0.06, container=doc)
    return p


def body_placeholder(doc, text, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, family=FONT_BODY, size=size, color=NEUTRAL_700)
    return p


def build():
    doc = Document()
    set_default_fonts(doc)
    section = doc.sections[0]
    set_margins(section, top=2.2, bottom=2.0, left=2.5, right=2.5)

    # ---------------- Cover ----------------
    for _ in range(3):
        doc.add_paragraph()
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_logo.add_run()
    r.add_picture(str(LOGO_ICON_NAVY), height=Cm(2.2))

    p_word = doc.add_paragraph()
    p_word.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_word.paragraph_format.space_before = Pt(8)
    r1 = p_word.add_run("alio")
    set_run_font(r1, family=FONT_DISPLAY, size=22, bold=True, color=NEUTRAL_950)
    r2 = p_word.add_run(" analytics")
    set_run_font(r2, family=FONT_DISPLAY, size=22, bold=False, color=NEUTRAL_700)

    for _ in range(2):
        doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run("Proposta")
    set_run_font(r, family=FONT_DISPLAY, size=44, bold=True, color=PRIMARY_700)

    p_project = doc.add_paragraph()
    p_project.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_project.paragraph_format.space_before = Pt(6)
    r = p_project.add_run("{{Nome do Projeto}}")
    set_run_font(r, family=FONT_DISPLAY, size=20, color=ACCENT_600)

    for _ in range(4):
        doc.add_paragraph()

    bar_p = doc.add_paragraph()
    bar_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_gradient_bar(doc, 4.5, height_cm=0.1, container=doc)

    for label, value in (
        ("Preparado para", "{{Nome do Cliente}}"),
        ("Data", "{{Data}}"),
        ("Validade da proposta", "{{Validade — ex: 30 dias}}"),
        ("Preparado por", f"{COMPANY_NAME}"),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        r = p.add_run(f"{label}:  ")
        set_run_font(r, family=FONT_BODY, size=11, bold=True, color=NEUTRAL_500)
        r2 = p.add_run(value)
        set_run_font(r2, family=FONT_BODY, size=11, color=NEUTRAL_950)

    add_page_break(doc)

    # ---------------- 1. About Alio ----------------
    section_heading(doc, 1, "Sobre a Alio Analytics")
    body_placeholder(
        doc,
        f"A {COMPANY_NAME} é uma empresa de tecnologia sediada em Luanda que transforma "
        "sistemas complexos em software claro e fiável — engenharia e design de "
        "classe mundial, construídos em Angola. {{Adicionar contexto específico do "
        "cliente/setor aqui.}}",
    )
    body_placeholder(doc, f"“{TAGLINE_EN}”")

    # ---------------- 2. Understanding / Scope ----------------
    section_heading(doc, 2, "Entendimento & Âmbito")
    body_placeholder(doc, "{{Descrever o problema do cliente, o contexto e os objetivos "
                           "do projeto tal como entendidos pela Alio.}}")
    for b in ["{{Objetivo 1}}", "{{Objetivo 2}}", "{{Objetivo 3}}"]:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Cm(0.6)
        r = p.add_run("•  " + b)
        set_run_font(r, family=FONT_BODY, size=11, color=NEUTRAL_700)

    # ---------------- 3. Approach ----------------
    section_heading(doc, 3, "Abordagem & Metodologia")
    body_placeholder(doc, "{{Descrever a metodologia de execução — fases, governação, "
                           "cadência de comunicação com o cliente.}}")

    # ---------------- 4. Deliverables ----------------
    section_heading(doc, 4, "Entregáveis")
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ("Entregável", "Descrição")):
        set_cell_background(cell, PRIMARY_700)
        cell.paragraphs[0].text = ""
        r = cell.paragraphs[0].add_run(text)
        set_run_font(r, family=FONT_DISPLAY, size=11, bold=True, color=WHITE)
    for _ in range(3):
        row = table.add_row().cells
        row[0].paragraphs[0].text = ""
        r = row[0].paragraphs[0].add_run("{{Entregável}}")
        set_run_font(r, family=FONT_BODY, size=10, color=NEUTRAL_950)
        row[1].paragraphs[0].text = ""
        r2 = row[1].paragraphs[0].add_run("{{Descrição breve}}")
        set_run_font(r2, family=FONT_BODY, size=10, color=NEUTRAL_700)

    # ---------------- 5. Timeline ----------------
    section_heading(doc, 5, "Cronograma")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ("Fase", "Duração", "Marco")):
        set_cell_background(cell, PRIMARY_700)
        cell.paragraphs[0].text = ""
        r = cell.paragraphs[0].add_run(text)
        set_run_font(r, family=FONT_DISPLAY, size=11, bold=True, color=WHITE)
    for phase in ["{{Fase 1}}", "{{Fase 2}}", "{{Fase 3}}"]:
        row = table.add_row().cells
        for cell, val in zip(row, (phase, "{{Duração}}", "{{Marco/Entrega}}")):
            cell.paragraphs[0].text = ""
            r = cell.paragraphs[0].add_run(val)
            set_run_font(r, family=FONT_BODY, size=10, color=NEUTRAL_700)

    # ---------------- 6. Pricing (Kwanza) ----------------
    section_heading(doc, 6, "Investimento")
    body_placeholder(doc, "Valores em Kwanza (Kz). Faturação por marcos, salvo indicação em "
                           "contrário. Preços não incluem IVA (14%), acrescido conforme a lei.")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ("Item / Marco", "Descrição", "Valor (Kz)")):
        set_cell_background(cell, PRIMARY_700)
        cell.paragraphs[0].text = ""
        r = cell.paragraphs[0].add_run(text)
        set_run_font(r, family=FONT_DISPLAY, size=11, bold=True, color=WHITE)
    for item in ["{{Marco 1}}", "{{Marco 2}}", "{{Marco 3}}"]:
        row = table.add_row().cells
        for cell, val in zip(row, (item, "{{Descrição}}", "{{Kz 0,00}}")):
            cell.paragraphs[0].text = ""
            r = cell.paragraphs[0].add_run(val)
            set_run_font(r, family=FONT_BODY, size=10, color=NEUTRAL_700)
    total_row = table.add_row().cells
    total_row[0].merge(total_row[1])
    total_row[0].paragraphs[0].text = ""
    r = total_row[0].paragraphs[0].add_run("Total (sem IVA)")
    set_run_font(r, family=FONT_DISPLAY, size=11, bold=True, color=NEUTRAL_950)
    total_row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    total_row[2].paragraphs[0].text = ""
    r2 = total_row[2].paragraphs[0].add_run("{{Kz 0,00}}")
    set_run_font(r2, family=FONT_MONO, size=11, bold=True, color=PRIMARY_700)

    # ---------------- 7. Terms ----------------
    section_heading(doc, 7, "Termos")
    for t in [
        "Condições de pagamento: {{ex. 30% adjudicação, 40% intermédio, 30% entrega}}.",
        "Prazo de validade da proposta: {{30 dias}} a partir da data de emissão.",
        "Alterações de âmbito estão sujeitas a aditamento e reavaliação de preço/prazo.",
        "Lei aplicável: legislação da República de Angola.",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run("•  " + t)
        set_run_font(r, family=FONT_BODY, size=10.5, color=NEUTRAL_700)

    # ---------------- 8. Signature / Acceptance ----------------
    add_page_break(doc)
    section_heading(doc, 8, "Assinatura & Aceitação")
    body_placeholder(doc, "A aceitação desta proposta implica a concordância com o âmbito, "
                           "prazos e condições de investimento aqui descritos.")

    sign_table = doc.add_table(rows=4, cols=2)
    strip_table_borders(sign_table)
    labels = [
        (f"Pela {COMPANY_NAME}", "Pelo Cliente"),
        ("{{Nome}}", "{{Nome}}"),
        ("{{Cargo}}", "{{Cargo}}"),
        ("Data: {{data}}", "Data: {{data}}"),
    ]
    for i, (left, right) in enumerate(labels):
        cells = sign_table.rows[i].cells
        for cell, val, bold in ((cells[0], left, i == 0), (cells[1], right, i == 0)):
            cell.paragraphs[0].text = ""
            r = cell.paragraphs[0].add_run(val)
            set_run_font(r, family=FONT_BODY, size=11, bold=bold, color=NEUTRAL_950 if bold else NEUTRAL_700)
    # signature line spacer row
    line_row = doc.add_table(rows=1, cols=2)
    strip_table_borders(line_row)
    for cell in line_row.rows[0].cells:
        cell.paragraphs[0].text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(30)
        r = p.add_run("_____________________________")
        set_run_font(r, family=FONT_BODY, size=11, color=NEUTRAL_200)

    # ---------------- Footer (fiscal, all pages) ----------------
    footer = section.footer
    footer.is_linked_to_previous = False
    add_gradient_bar(doc, PAGE_WIDTH_CM, height_cm=0.06, container=footer)
    add_placeholder_line(
        footer, f"{COMPANY_NAME} · NIF {NIF} · {ADDRESS}",
        size=8, color=NEUTRAL_500, align=WD_ALIGN_PARAGRAPH.CENTER, family=FONT_MONO, space_after=1,
    )
    add_placeholder_line(
        footer, f"IBAN {IBAN} · {BANK_NAME} · {WEBSITE} · {EMAIL}",
        size=8, color=NEUTRAL_500, align=WD_ALIGN_PARAGRAPH.CENTER, family=FONT_MONO, space_after=0,
    )

    doc.core_properties.title = "Alio Analytics — Proposta"
    doc.core_properties.author = "Alio Analytics"

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    docx_path = OUT_DIR / "alio-proposal.docx"
    dotx_path = OUT_DIR / "alio-proposal.dotx"
    doc.save(docx_path)
    docx_to_dotx(docx_path, dotx_path)
    print("  OK", docx_path.name, "+", dotx_path.name)


if __name__ == "__main__":
    build()
