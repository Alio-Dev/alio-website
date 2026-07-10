"""Shared helpers for the Alio Word templates (Letterhead, Proposal)."""

from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from brand import (
    FONT_DISPLAY, FONT_BODY, FONT_MONO,
    PRIMARY_700, ACCENT_500, NEUTRAL_950, NEUTRAL_700, NEUTRAL_500, NEUTRAL_200, WHITE,
    hx,
)

RGB = lambda h: RGBColor(*hx(h))


def set_default_fonts(document):
    """Normal style -> Instrument Sans (with a safe system fallback chain)."""
    style = document.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = Pt(11)
    style.font.color.rgb = RGB(NEUTRAL_950)
    # East-Asian/complex-script slots must also be set or Word substitutes Calibri.
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), FONT_BODY)


def set_run_font(run, family=FONT_BODY, size=11, bold=False, color=NEUTRAL_950, italic=False):
    run.font.name = family
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGB(color)
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), family)


def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def strip_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)


def add_gradient_bar(document, width_cm, height_cm=0.12, container=None):
    """
    Word has no native gradient fill for a simple rule, so the brand gradient
    (cyan -> indigo, BRAND.md sec.3) is approximated with a thin two-cell,
    borderless table: left cell accent-500, right cell primary-700.
    """
    host = container if container is not None else document
    total_width = Cm(width_cm)
    try:
        table = host.add_table(rows=1, cols=2, width=total_width)
    except TypeError:
        # python-docx's Document.add_table() takes no width kwarg (it is set
        # per-column below); Header/Footer's add_table() requires it.
        table = host.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    strip_table_borders(table)
    half = Cm(width_cm / 2)
    table.columns[0].width = half
    table.columns[1].width = half
    row = table.rows[0]
    row.height = Cm(height_cm)
    for cell, color in zip(table.rows[0].cells, (ACCENT_500, PRIMARY_700)):
        cell.width = half
        set_cell_background(cell, color)
        for p in cell.paragraphs:
            p.text = ""
    return table


def add_picture_paragraph(document, image_path, width_cm, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                           container=None, space_after=0):
    host = container if container is not None else document
    p = host.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    return p


def add_placeholder_line(container, text, size=10, color=NEUTRAL_500, align=WD_ALIGN_PARAGRAPH.LEFT,
                          bold=False, family=FONT_BODY, space_after=2):
    p = container.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_run_font(r, family=family, size=size, bold=bold, color=color)
    return p


def set_margins(section, top=2.5, bottom=2.2, left=2.5, right=2.5):
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)
