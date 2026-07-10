"""
Generate Letterhead.docx (+ Letterhead.dotx template) — Alio Analytics.

FE-1 brand-fidelity checklist applied:
  - Logo (navy icon) + wordmark in the header, brand gradient rule (BRAND.md
    sec.3 "Gradient: linear-gradient(90deg,#07B7D1,#2B3990)") approximated as
    a two-tone bar (Word has no native line-gradient — see docx_helpers).
  - Headings/salutation in Sora, body in Instrument Sans, fiscal footer in
    JetBrains Mono (BRAND.md sec.4 "Code/data/labels: JetBrains Mono").
  - Fiscal footer fields exactly as BRAND.md sec.10.
"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from brand import (
    FONT_DISPLAY, FONT_BODY, FONT_MONO,
    PRIMARY_700, NEUTRAL_950, NEUTRAL_700, NEUTRAL_500,
    COMPANY_NAME, NIF, ADDRESS, IBAN, BANK_NAME, EMAIL, WEBSITE, PHONE,
    LOGO_ICON_NAVY, OUT_DIR,
)
from docx_helpers import (
    set_default_fonts, set_run_font, add_gradient_bar, add_placeholder_line,
    set_margins,
)
from to_template import docx_to_dotx

PAGE_WIDTH_CM = 21.0 - 2.5 - 2.5  # A4 minus L/R margins


def build():
    doc = Document()
    set_default_fonts(doc)
    section = doc.sections[0]
    set_margins(section, top=2.0, bottom=1.8, left=2.5, right=2.5)

    header = section.header
    header.is_linked_to_previous = False
    htable_container = header

    # --- Header: logo + wordmark, then the brand gradient rule ---
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(2)
    run_icon = hp.add_run()
    run_icon.add_picture(str(LOGO_ICON_NAVY), height=Cm(1.0))
    run_gap = hp.add_run("   ")
    run_alio = hp.add_run("alio")
    set_run_font(run_alio, family=FONT_DISPLAY, size=20, bold=True, color=NEUTRAL_950)
    run_ana = hp.add_run(" analytics")
    set_run_font(run_ana, family=FONT_DISPLAY, size=20, bold=False, color=NEUTRAL_700)

    add_gradient_bar(doc, PAGE_WIDTH_CM, height_cm=0.1, container=header)

    # --- Footer: fiscal details (mono, per BRAND.md sec.10) ---
    footer = section.footer
    footer.is_linked_to_previous = False
    add_gradient_bar(doc, PAGE_WIDTH_CM, height_cm=0.06, container=footer)
    add_placeholder_line(
        footer, f"{COMPANY_NAME} · NIF {NIF} · {ADDRESS}",
        size=8, color=NEUTRAL_500, align=WD_ALIGN_PARAGRAPH.CENTER, family=FONT_MONO, space_after=1,
    )
    add_placeholder_line(
        footer, f"IBAN {IBAN} · {BANK_NAME} · {WEBSITE} · {EMAIL} · {PHONE}",
        size=8, color=NEUTRAL_500, align=WD_ALIGN_PARAGRAPH.CENTER, family=FONT_MONO, space_after=0,
    )

    # --- Body: date, salutation, empty writing area, closing ---
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_date.paragraph_format.space_before = Pt(18)
    r = p_date.add_run("Luanda, {{Data}}")
    set_run_font(r, family=FONT_BODY, size=11, color=NEUTRAL_700, italic=True)

    p_subject = doc.add_paragraph()
    p_subject.paragraph_format.space_before = Pt(12)
    r = p_subject.add_run("Assunto: {{Assunto da carta}}")
    set_run_font(r, family=FONT_DISPLAY, size=12, bold=True, color=PRIMARY_700)

    p_recipient = doc.add_paragraph()
    p_recipient.paragraph_format.space_before = Pt(6)
    r = p_recipient.add_run("{{Nome / Empresa do Destinatário}}\n{{Morada do Destinatário}}")
    set_run_font(r, family=FONT_BODY, size=11, color=NEUTRAL_700)

    p_greeting = doc.add_paragraph()
    p_greeting.paragraph_format.space_before = Pt(18)
    r = p_greeting.add_run("Exmo.(a) Senhor(a),")
    set_run_font(r, family=FONT_BODY, size=11, bold=True, color=NEUTRAL_950)

    # Empty writing area (the deliverable: a blank body to type the letter).
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        set_default_fonts(doc)  # no-op safeguard; Normal style already set

    p_closing = doc.add_paragraph()
    p_closing.paragraph_format.space_before = Pt(24)
    r = p_closing.add_run("Com os melhores cumprimentos,")
    set_run_font(r, family=FONT_BODY, size=11, color=NEUTRAL_950)

    for _ in range(2):
        doc.add_paragraph()

    p_sign = doc.add_paragraph()
    r = p_sign.add_run("{{Nome do Signatário}}")
    set_run_font(r, family=FONT_BODY, size=11, bold=True, color=NEUTRAL_950)
    p_role = doc.add_paragraph()
    r = p_role.add_run("{{Cargo}} · Alio Analytics")
    set_run_font(r, family=FONT_BODY, size=10, color=NEUTRAL_500)

    doc.core_properties.title = "Alio Analytics — Letterhead"
    doc.core_properties.author = "Alio Analytics"

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    docx_path = OUT_DIR / "alio-letterhead.docx"
    dotx_path = OUT_DIR / "alio-letterhead.dotx"
    doc.save(docx_path)
    docx_to_dotx(docx_path, dotx_path)
    print("  OK", docx_path.name, "+", dotx_path.name)


if __name__ == "__main__":
    build()
